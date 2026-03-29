import os
from crewai.tools import BaseTool
from pydantic import Field


class ResumeParserTool(BaseTool):
    """Parses a resume file (PDF or DOCX) and extracts its text content."""

    name: str = "resume_parser"
    description: str = (
        "Extracts text content from a resume file. "
        "Returns the raw text content of the resume for further structured extraction. "
        "Call this tool with any input (e.g. 'parse') to extract the pre-configured resume."
    )
    resume_path: str = Field(default="", description="Path to the resume file to parse")

    def _run(self, file_path: str = "") -> str:
        # Use pre-configured path as primary, fall back to argument
        path = self.resume_path or file_path
        if not path:
            return "Error: No file path provided. Please specify a resume file path."

        # Clean up path - agent may wrap it in quotes
        path = path.strip().strip("'\"")

        if not os.path.exists(path):
            return f"Error: File not found at '{path}'. Please check the path."

        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(path)
        elif ext == ".txt":
            return self._parse_text(path)
        else:
            return f"Error: Unsupported file format '{ext}'. Supported formats: .pdf, .docx, .txt"

    def _parse_pdf(self, path: str) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            if not text_parts:
                return "Error: Could not extract text from PDF. The file may be image-based or empty."

            return "\n\n".join(text_parts)
        except ImportError:
            return "Error: pypdf is not installed. Run: pip install pypdf"
        except Exception as e:
            return f"Error parsing PDF: {str(e)}"

    def _parse_docx(self, path: str) -> str:
        try:
            from docx import Document

            doc = Document(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            if not paragraphs:
                return "Error: Could not extract text from DOCX. The file may be empty."

            return "\n".join(paragraphs)
        except ImportError:
            return "Error: python-docx is not installed. Run: pip install python-docx"
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"

    def _parse_text(self, path: str) -> str:
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            return f"Error reading text file: {str(e)}"
